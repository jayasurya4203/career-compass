import io
import re
from typing import Any

import pdfplumber
from docx import Document

SKILL_LEXICON = [
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "node",
    "sql",
    "html",
    "css",
    "django",
    "flask",
    "fastapi",
    "pandas",
    "numpy",
    "machine learning",
    "deep learning",
    "tensorflow",
    "pytorch",
    "aws",
    "azure",
    "docker",
    "kubernetes",
    "git",
    "linux",
    "excel",
    "tableau",
    "power bi",
    "figma",
    "kotlin",
    "android",
    "selenium",
    "terraform",
    "kubernetes",
    "rest",
    "api",
    "mongodb",
    "postgresql",
    "mysql",
    "c++",
    "c#",
    "go",
    "rust",
    "spring",
    "angular",
    "vue",
]


def extract_text_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            parts.append(t)
    return "\n".join(parts)


def extract_text_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Unsupported file type. Use PDF or DOCX.")


def parse_resume_text(text: str) -> dict[str, Any]:
    low = text.lower()
    email_m = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    phone_m = re.search(r"(\+?\d[\d\s\-]{8,}\d)", text)
    gh = re.search(r"github\.com/[\w\-]+/?", low)
    li = re.search(r"linkedin\.com/in/[\w\-]+/?", low)
    skills_found = []
    for s in SKILL_LEXICON:
        if s in low:
            skills_found.append(s.title() if s != "c++" else "C++")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    name_guess = lines[0][:80] if lines else ""
    edu = []
    for ln in lines:
        if any(k in ln.lower() for k in ("university", "college", "institute", "b.tech", "b.e", "m.tech", "cgpa", "gpa")):
            edu.append(ln[:200])
    projects = []
    if "project" in low:
        for ln in lines:
            if "project" in ln.lower():
                projects.append(ln[:200])
    return {
        "extracted_name": name_guess,
        "extracted_email": email_m.group(0) if email_m else "",
        "extracted_phone": phone_m.group(0).strip() if phone_m else "",
        "extracted_skills": ", ".join(sorted(set(skills_found))),
        "extracted_projects": " | ".join(projects[:5]),
        "extracted_education": " | ".join(edu[:5]),
        "extracted_links": ", ".join(filter(None, [gh.group(0) if gh else "", li.group(0) if li else ""])),
        "resume_completeness_hint": _completeness(low),
    }


def _completeness(low: str) -> float:
    keys = ["education", "skill", "project", "experience", "intern", "certif"]
    return round(sum(1 for k in keys if k in low) / len(keys), 2)
